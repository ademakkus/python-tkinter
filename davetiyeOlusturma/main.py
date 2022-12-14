import sys
from docx import Documnet

dosya=Document()
yer=input('Yer Bilgisi:')
vakit=input('Tarih ve saat:')
liste=sys.stdin.read().split('\n')
liste=list(map(str.strip,liste))
for i in range(len(liste)):
    dosya.add_heading('Saygı değer '+liste[i]+'!',0)
    dosya.add_picture('ring.jpg')
dosya.save('davetiye.docx')